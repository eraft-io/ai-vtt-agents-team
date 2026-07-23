"""Markdown 转 DOCX 转换工具。

将包含图片引用的 Markdown 文本转换为格式化的 Word 文档，
支持标题、段落、加粗、斜体和嵌入图片。
"""

import logging
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Markdown 图片正则: ![alt](path)
_IMG_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
# 行内样式
_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_PATTERN = re.compile(r"\*(.+?)\*")


def _apply_inline_styles(paragraph, text: str) -> None:
    """为段落文本添加加粗和斜体的行内样式。"""
    # 拆分 token: **bold**, *italic*, plain text
    # 用一个简单的正则分词
    token_re = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    parts = token_re.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _resolve_image_path(src: str, project_dir: str) -> str | None:
    """解析图片路径，支持绝对路径和相对路径（相对于 project_dir）。"""
    if os.path.isabs(src) and os.path.isfile(src):
        return src
    # 尝试相对于项目目录
    rel = os.path.join(project_dir, src)
    if os.path.isfile(rel):
        return rel
    # 尝试去掉 leading /
    stripped = src.lstrip("/")
    rel2 = os.path.join(project_dir, stripped)
    if os.path.isfile(rel2):
        return rel2
    return None


def markdown_to_docx(
    md_text: str,
    output_path: str,
    project_dir: str = "",
    title: str = "",
) -> str:
    """将 Markdown 文本转换为 DOCX 文件。

    Args:
        md_text: Markdown 格式的文本内容。
        output_path: 输出的 .docx 文件路径。
        project_dir: 项目目录（用于解析图片的相对路径）。
        title: 文档标题（可选，不传则从第一个 # 标题取）。

    Returns:
        实际写入的 .docx 文件路径。
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    has_title = False

    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 跳过分隔线
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # 标题 (# ## ###)
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if not has_title and level == 1 and not title:
                # 第一个 H1 作为文档标题
                doc.add_heading(heading_text, level=0)
                has_title = True
            else:
                doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # 图片行（整行是图片）
        img_match = _IMG_PATTERN.match(line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_src = img_match.group(2)
            resolved = _resolve_image_path(img_src, project_dir)
            if resolved:
                try:
                    doc.add_picture(resolved, width=Inches(5.5))
                    # 图片居中
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt_text:
                        caption = doc.add_paragraph(alt_text)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].font.size = Pt(9)
                        caption.runs[0].font.italic = True
                except Exception as e:
                    logger.warning("无法插入图片 %s: %s", resolved, e)
                    doc.add_paragraph(f"[图片: {alt_text or img_src}]")
            else:
                doc.add_paragraph(f"[图片: {alt_text or img_src}]")
            i += 1
            continue

        # 列表项 (- * +)
        list_match = re.match(r"^[\-\*\+]\s+(.+)$", line.strip())
        if list_match:
            list_text = list_match.group(1)
            para = doc.add_paragraph(style="List Bullet")
            _apply_inline_styles(para, list_text)
            i += 1
            continue

        # 有序列表 (1. 2. etc.)
        ol_match = re.match(r"^\d+\.\s+(.+)$", line.strip())
        if ol_match:
            ol_text = ol_match.group(1)
            para = doc.add_paragraph(style="List Number")
            _apply_inline_styles(para, ol_text)
            i += 1
            continue

        # 代码块 ```
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(code_text)
            code_run.font.name = "Courier New"
            code_run.font.size = Pt(9)
            continue

        # 普通段落（可能包含行内图片）
        # 先检查段落中是否有图片
        if _IMG_PATTERN.search(line):
            # 把行拆分为文本和图片交替
            _add_mixed_paragraph(doc, line, project_dir)
        else:
            para = doc.add_paragraph()
            _apply_inline_styles(para, line)

        i += 1

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("DOCX 已保存: %s", output_path)
    return output_path


def _add_mixed_paragraph(doc: Document, line: str, project_dir: str) -> None:
    """处理包含行内图片的段落，文本和图片交替插入。"""
    # 按图片分割
    parts = _IMG_PATTERN.split(line)
    # parts 格式: [text, alt, src, text, alt, src, ...]
    i = 0
    while i < len(parts):
        if i + 2 < len(parts):
            # text + img
            text = parts[i]
            alt = parts[i + 1]
            src = parts[i + 2]
            if text.strip():
                para = doc.add_paragraph()
                _apply_inline_styles(para, text)
            resolved = _resolve_image_path(src, project_dir)
            if resolved:
                try:
                    doc.add_picture(resolved, width=Inches(5.5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"[图片: {alt or src}]")
            else:
                doc.add_paragraph(f"[图片: {alt or src}]")
            i += 3
        else:
            # 剩余文本
            text = parts[i]
            if text.strip():
                para = doc.add_paragraph()
                _apply_inline_styles(para, text)
            i += 1
